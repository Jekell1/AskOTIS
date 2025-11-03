#!/usr/bin/env python3
"""
Export Chapter 5: Setting Modernization Goals to Word document
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re

def add_heading(doc, text, level=1):
    """Add a styled heading to the document."""
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, style='Normal'):
    """Add a paragraph with specified style."""
    p = doc.add_paragraph(text, style=style)
    return p

def add_code_block(doc, code_text):
    """Add a code block with monospace font."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_table_from_markdown(doc, markdown_table):
    """Convert a markdown table to a Word table."""
    lines = [line.strip() for line in markdown_table.strip().split('\n') if line.strip()]
    if len(lines) < 2:
        return None
    
    # Parse header
    headers = [cell.strip() for cell in lines[0].split('|') if cell.strip()]
    num_cols = len(headers)
    
    # Parse data rows (skip separator line)
    data_rows = []
    for line in lines[2:]:  # Skip header and separator
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if len(cells) == num_cols:
            data_rows.append(cells)
    
    if not data_rows:
        return None
    
    # Create table
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add data
    for i, row_data in enumerate(data_rows):
        row_cells = table.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    return table

def parse_and_add_markdown_content(doc, markdown_text):
    """Parse markdown content and add to Word document."""
    lines = markdown_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            add_heading(doc, line[2:], level=1)
        elif line.startswith('## '):
            add_heading(doc, line[3:], level=2)
        elif line.startswith('### '):
            add_heading(doc, line[4:], level=3)
        elif line.startswith('#### '):
            add_heading(doc, line[5:], level=4)
        
        # Code blocks
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, '\n'.join(code_lines))
        
        # Tables
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            add_table_from_markdown(doc, '\n'.join(table_lines))
            i -= 1
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            add_paragraph(doc, line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            add_paragraph(doc, text, style='List Number')
        
        # Bold text
        elif line.startswith('**') and line.endswith('**'):
            p = add_paragraph(doc, line.strip('*'))
            for run in p.runs:
                run.font.bold = True
        
        # Regular paragraph
        else:
            add_paragraph(doc, line)
        
        i += 1

def main():
    """Convert Chapter 5 markdown to Word document."""
    
    # Read the markdown file
    with open("Chapter5_Modernization_Goals.md", "r", encoding="utf-8") as f:
        markdown_content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set document title
    title = doc.add_heading('Chapter 5: Setting Modernization Goals', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Parse and add content
    parse_and_add_markdown_content(doc, markdown_content)
    
    # Save document
    output_file = "Chapter5_Modernization_Goals.docx"
    doc.save(output_file)
    print(f"✅ Word document created: {output_file}")

if __name__ == "__main__":
    main()
