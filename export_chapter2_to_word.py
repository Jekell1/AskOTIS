"""
Export Chapter 2 Markdown to Word Document
Converts the Chapter 2 business domain analysis to a professionally formatted Word document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re

def create_styled_document():
    """Create a new document with custom styles"""
    doc = Document()
    
    # Configure styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_font = title_style.font
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = RGBColor(31, 56, 100)
    
    # Heading 1 style  
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.size = Pt(18)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(31, 56, 100)
    
    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(68, 114, 196)
    
    # Heading 3 style
    h3_style = styles['Heading 3']
    h3_font = h3_style.font
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(68, 114, 196)
    
    return doc

def parse_markdown_line(line):
    """Parse a markdown line and return its type and content"""
    line = line.rstrip()
    
    if not line:
        return 'empty', ''
    
    # Check for headers
    if line.startswith('# '):
        return 'title', line[2:]
    elif line.startswith('## '):
        return 'h1', line[3:]
    elif line.startswith('### '):
        return 'h2', line[4:]
    elif line.startswith('#### '):
        return 'h3', line[5:]
    
    # Check for bullet lists
    if line.startswith('- ') or line.startswith('* '):
        return 'bullet', line[2:]
    
    # Check for bold text
    if line.startswith('**') and line.endswith('**'):
        return 'bold', line[2:-2]
    
    # Check for code blocks
    if line.startswith('```'):
        return 'code_marker', line[3:]
    
    # Check for horizontal rules
    if line.startswith('---'):
        return 'hr', ''
    
    # Regular paragraph
    return 'paragraph', line

def add_formatted_text(paragraph, text):
    """Add text with inline formatting (bold, italic, code)"""
    # Pattern to match **bold**, *italic*, and `code`
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith('*') and part.endswith('*'):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        elif part.startswith('`') and part.endswith('`'):
            # Code text
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        else:
            # Regular text
            paragraph.add_run(part)

def convert_markdown_to_word(md_file, docx_file):
    """Convert markdown file to Word document"""
    doc = create_styled_document()
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        line_type, content = parse_markdown_line(line)
        
        # Handle code blocks
        if line_type == 'code_marker':
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # End of code block
                in_code_block = False
                if code_lines:
                    code_para = doc.add_paragraph()
                    code_para.style = 'Normal'
                    code_run = code_para.add_run('\n'.join(code_lines))
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(9)
                    code_para.paragraph_format.left_indent = Inches(0.5)
                    code_para.paragraph_format.space_before = Pt(6)
                    code_para.paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line.rstrip())
            i += 1
            continue
        
        # Handle table detection
        if '|' in line and not in_table:
            # Start of table
            in_table = True
            table_lines = [line]
            i += 1
            continue
        
        if in_table:
            if '|' in line:
                table_lines.append(line)
                i += 1
                continue
            else:
                # End of table
                in_table = False
                # Process table
                process_table(doc, table_lines)
                table_lines = []
                # Don't increment i, process current line normally
        
        # Handle different line types
        if line_type == 'title':
            p = doc.add_heading(content, level=0)
        elif line_type == 'h1':
            doc.add_heading(content, level=1)
        elif line_type == 'h2':
            doc.add_heading(content, level=2)
        elif line_type == 'h3':
            doc.add_heading(content, level=3)
        elif line_type == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, content)
        elif line_type == 'bold':
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.bold = True
        elif line_type == 'hr':
            doc.add_paragraph('_' * 50)
        elif line_type == 'empty':
            pass  # Skip empty lines
        else:  # paragraph
            p = doc.add_paragraph()
            add_formatted_text(p, content)
        
        i += 1
    
    # Handle any remaining table
    if in_table and table_lines:
        process_table(doc, table_lines)
    
    # Save document
    doc.save(docx_file)
    print(f"Document saved: {docx_file}")

def process_table(doc, table_lines):
    """Process and add a table to the document"""
    if len(table_lines) < 2:
        return
    
    # Parse header
    header_cells = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
    
    # Skip separator line
    if len(table_lines) < 3:
        return
    
    # Parse data rows
    data_rows = []
    for line in table_lines[2:]:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            data_rows.append(cells)
    
    if not data_rows:
        return
    
    # Create table
    table = doc.add_table(rows=1 + len(data_rows), cols=len(header_cells))
    table.style = 'Light Grid Accent 1'
    
    # Add header with bold formatting
    header_row = table.rows[0]
    for i, cell_text in enumerate(header_cells):
        cell = header_row.cells[i]
        cell.text = cell_text
        # Make header bold
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add data rows
    for row_idx, row_data in enumerate(data_rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(row.cells):
                row.cells[col_idx].text = cell_text
    
    doc.add_paragraph()  # Add space after table

def main():
    md_file = 'Chapter2_Business_Domain.md'
    docx_file = 'Chapter2_Business_Domain.docx'
    
    print(f"Converting {md_file} to {docx_file}...")
    convert_markdown_to_word(md_file, docx_file)
    print("Conversion complete!")

if __name__ == '__main__':
    main()
