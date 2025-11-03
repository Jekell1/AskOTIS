"""
Export Chapter 3 markdown to Word document format.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def create_chapter_word_doc():
    """Convert Chapter 3 markdown to formatted Word document."""
    
    # Read the markdown content
    with open('Chapter3_Technical_Architecture.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Main title (single #)
        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
            
        # Chapter heading (## Chapter)
        elif line.startswith('## Chapter'):
            heading = line[3:].strip()
            p = doc.add_heading(heading, level=1)
            run = p.runs[0]
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 51, 102)
            
        # Section heading (###)
        elif line.startswith('### '):
            heading = line[4:].strip()
            p = doc.add_heading(heading, level=2)
            run = p.runs[0]
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 102, 204)
            
        # Subsection heading (####)
        elif line.startswith('#### '):
            heading = line[5:].strip()
            p = doc.add_heading(heading, level=3)
            run = p.runs[0]
            run.font.size = Pt(12)
            run.font.bold = True
            
        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph()  # Add spacing
            
        # Code block
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
        
        # Table rows (simplified table handling)
        elif line.startswith('|') and '|' in line[1:]:
            # Collect table rows
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # Back up one
            
            # Parse table
            if len(table_lines) > 2:  # Header + separator + data
                # Extract headers
                headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
                
                # Create table
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'
                
                # Add headers
                hdr_cells = table.rows[0].cells
                for idx, header in enumerate(headers):
                    hdr_cells[idx].text = header
                    for paragraph in hdr_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Add data rows (skip separator row)
                for row_line in table_lines[2:]:
                    cells_data = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    row_cells = table.add_row().cells
                    for idx, cell_data in enumerate(cells_data):
                        row_cells[idx].text = cell_data
        
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Handle bold text **text**
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            # Handle code `text`
            text = re.sub(r'`(.*?)`', r'\1', text)
            p = doc.add_paragraph(text, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(11)
            
        # Regular paragraph
        else:
            # Handle bold text **text** and code `text`
            if '**' in line or '`' in line:
                p = doc.add_paragraph()
                # Split by code and bold markers
                parts = re.split(r'(\*\*.*?\*\*|`.*?`)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                        run.font.size = Pt(11)
                    elif part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
                    else:
                        run = p.add_run(part)
                        run.font.size = Pt(11)
            else:
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.size = Pt(11)
            
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(6)
        
        i += 1
    
    # Save document
    output_file = 'Chapter3_Technical_Architecture.docx'
    doc.save(output_file)
    print(f"✅ Word document created: {output_file}")
    return output_file

if __name__ == '__main__':
    create_chapter_word_doc()
