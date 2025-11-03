#!/usr/bin/env python3
"""
Export Chapter 1: Introduction to Word format
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def add_heading(doc, text, level):
    """Add a heading with appropriate formatting."""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.size = Pt(24)
            run.font.bold = True
    elif level == 2:
        for run in heading.runs:
            run.font.size = Pt(18)
            run.font.bold = True
    elif level == 3:
        for run in heading.runs:
            run.font.size = Pt(14)
            run.font.bold = True

def process_table(doc, table_lines):
    """Process and add a table to the document."""
    if len(table_lines) < 2:
        return
    
    # Parse header
    header_cells = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
    
    # Skip separator line
    if len(table_lines) < 3:
        return
    
    # Parse data rows
    data_rows = []
    for line in table_lines[2:]:
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells and len(cells) == len(header_cells):
            data_rows.append(cells)
    
    if not data_rows:
        return
    
    # Create table
    table = doc.add_table(rows=1 + len(data_rows), cols=len(header_cells))
    table.style = 'Light Grid Accent 1'
    
    # Add header with bold formatting
    header_row = table.rows[0]
    for i, cell_text in enumerate(header_cells):
        cell = header_row.cells[i]
        cell.text = cell_text
        # Make header bold
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Add data rows
    for row_idx, row_data in enumerate(data_rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            row.cells[col_idx].text = cell_text

def add_paragraph(doc, text):
    """Add a paragraph with proper formatting."""
    if not text.strip():
        return
    
    # Handle code blocks
    if '```' in text:
        parts = re.split(r'```\w*\n?', text)
        for i, part in enumerate(parts):
            if i % 2 == 0:  # Regular text
                if part.strip():
                    p = doc.add_paragraph(part.strip())
                    p.paragraph_format.space_after = Pt(6)
            else:  # Code block
                p = doc.add_paragraph(part.strip())
                p.paragraph_format.left_indent = Inches(0.5)
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
    else:
        p = doc.add_paragraph(text.strip())
        p.paragraph_format.space_after = Pt(6)
        
        # Handle inline code
        if '`' in text:
            for run in p.runs:
                if '`' in run.text:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)

def main():
    # Read the markdown file
    with open('Chapter1_Introduction.md', 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Process markdown content
    i = 0
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Detect table start
        if '|' in line and not in_table:
            in_table = True
            table_lines = [line]
            i += 1
            continue
        
        # Collect table lines
        if in_table:
            if '|' in line:
                table_lines.append(line)
                i += 1
                continue
            else:
                # End of table
                process_table(doc, table_lines)
                table_lines = []
                in_table = False
                # Don't increment i, process current line normally
        
        if line.startswith('# '):
            add_heading(doc, line[2:], 1)
        elif line.startswith('## '):
            add_heading(doc, line[3:], 2)
        elif line.startswith('### '):
            add_heading(doc, line[4:], 3)
        elif line.startswith('#### '):
            add_heading(doc, line[5:], 4)
        elif line.strip():
            # Accumulate paragraph text
            para_text = line
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and '|' not in lines[i]:
                para_text += '\n' + lines[i].rstrip()
                i += 1
            add_paragraph(doc, para_text)
            continue
        
        i += 1
    
    # Handle any remaining table
    if in_table and table_lines:
        process_table(doc, table_lines)
    
    # Save document
    output_file = "Chapter1_Introduction.docx"
    doc.save(output_file)
    print(f"✅ Word document created: {output_file}")

if __name__ == "__main__":
    main()
