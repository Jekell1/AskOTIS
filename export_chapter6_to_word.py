"""
Export Chapter 6 to Word Document
Converts markdown to professionally formatted Word document
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def parse_markdown_to_word(md_file, output_file='Chapter6_Tools_Technologies.docx'):
    """Convert Chapter 6 markdown to Word document"""
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Read markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                code_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Title (# Chapter)
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Major sections (## Section)
        elif line.startswith('## '):
            heading = line[3:].strip()
            doc.add_heading(heading, level=1)
        
        # Subsections (### )
        elif line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, level=2)
        
        # Horizontal rules
        elif line.strip() == '---':
            doc.add_paragraph()  # Just add space
        
        # Bold text at start of line (like **Question N:**)
        elif line.startswith('**') and '**' in line[2:]:
            p = doc.add_paragraph()
            
            # Parse inline formatting
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    p.add_run(part)
        
        # Bullet lists
        elif line.startswith('- '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).strip()
            p = doc.add_paragraph(text, style='List Number')
        
        # Regular paragraphs
        elif line.strip():
            # Check for inline formatting
            if '**' in line or '`' in line:
                p = doc.add_paragraph()
                
                # Handle bold
                parts = re.split(r'(\*\*.*?\*\*|`.*?`)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                    elif part:
                        p.add_run(part)
            else:
                doc.add_paragraph(line)
        
        # Empty lines
        else:
            # Only add paragraph break if not after heading
            if i > 0 and not lines[i-1].startswith('#'):
                pass  # Paragraph spacing is automatic
        
        i += 1
    
    # Save document
    doc.save(output_file)
    print(f"✅ Word document created: {output_file}")

def main():
    parse_markdown_to_word('Chapter6_Tools_Technologies.md')

if __name__ == "__main__":
    main()
