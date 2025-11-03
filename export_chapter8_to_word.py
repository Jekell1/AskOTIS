"""
Export Chapter 8: Automated and Manual Code Analysis to Word format
Converts markdown to a formatted Word document using python-docx.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def parse_markdown_to_word(md_file, docx_file):
    """Convert markdown file to Word document."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Chapter title (# )
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Major sections (## )
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
            
        # Subsections (### )
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
            
        # Code blocks (```)
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines))
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 128)
        
        # Bold text (**text**)
        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        # Bullet points (- )
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        
        # Numbered lists (1. )
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line)
            doc.add_paragraph(text, style='List Number')
        
        # Regular paragraph
        elif line.strip():
            doc.add_paragraph(line)
        
        i += 1
    
    doc.save(docx_file)

def main():
    """Main execution."""
    md_file = "Chapter8_Code_Analysis.md"
    docx_file = "Chapter8_Code_Analysis.docx"
    
    parse_markdown_to_word(md_file, docx_file)
    print(f"✅ Word document created: {docx_file}")

if __name__ == "__main__":
    main()
