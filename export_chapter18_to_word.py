#!/usr/bin/env python3
"""
Export Chapter 18 to Word format with formatting
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

# Read markdown
with open('Chapter18_System_Maintenance.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Create Word document
doc = Document()

# Process content line by line
lines = content.split('\n')
i = 0
while i < len(lines):
    line = lines[i].strip()
    
    # Chapter title (# Title)
    if line.startswith('# Chapter'):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run(line[2:].strip())
        run.font.size = Pt(24)
        run.font.bold = True
        doc.add_paragraph()
    
    # Main sections (## Section)
    elif line.startswith('## Section') or line.startswith('## Introduction') or line.startswith('## Conclusion'):
        p = doc.add_paragraph()
        run = p.add_run(line[3:].strip())
        run.font.size = Pt(16)
        run.font.bold = True
        doc.add_paragraph()
    
    # Subsections (### Subsection)
    elif line.startswith('###'):
        p = doc.add_paragraph()
        run = p.add_run(line[4:].strip())
        run.font.size = Pt(14)
        run.font.bold = True
        doc.add_paragraph()
    
    # Horizontal rules
    elif line.startswith('---'):
        doc.add_paragraph()
    
    # Code blocks
    elif line.startswith('```'):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        code_text = '\n'.join(code_lines)
        p = doc.add_paragraph(code_text)
        p.style = 'Code'
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        doc.add_paragraph()
    
    # Bullet lists
    elif line.startswith('- ') or line.startswith('* '):
        text = line[2:].strip()
        # Handle bold in bullet points
        p = doc.add_paragraph(style='List Bullet')
        if '**' in text:
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    p.add_run(part)
        else:
            p.add_run(text)
    
    # Regular paragraphs
    elif line:
        p = doc.add_paragraph()
        # Handle bold text (**text**)
        if '**' in line:
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part:
                    p.add_run(part)
        # Handle inline code (`code`)
        elif '`' in line:
            parts = re.split(r'(`[^`]+`)', line)
            for part in parts:
                if part.startswith('`') and part.endswith('`'):
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                elif part:
                    p.add_run(part)
        else:
            p.add_run(line)
    
    i += 1

# Save document
doc.save('Chapter18_System_Maintenance.docx')
print("✅ Word document created: Chapter18_System_Maintenance.docx")
