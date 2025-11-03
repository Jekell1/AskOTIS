#!/usr/bin/env python3
"""
Export Chapter 13 to Word format
Converts Chapter13_Data_Migration.md to a formatted Word document.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def main():
    # Read markdown file
    with open("Chapter13_Data_Migration.md", "r", encoding="utf-8") as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Parse markdown and convert to Word
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Chapter title (# heading)
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(24)
            p.runs[0].font.bold = True
        
        # Section heading (## heading)
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:])
            p.runs[0].font.size = Pt(16)
            p.runs[0].font.bold = True
        
        # Subsection heading (### heading)
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:])
            p.runs[0].font.size = Pt(14)
            p.runs[0].font.bold = True
        
        # Code blocks (```)
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            p = doc.add_paragraph('\n'.join(code_lines))
            p.runs[0].font.name = 'Consolas'
            p.runs[0].font.size = Pt(9)
        
        # Bold questions (**Q#: text**)
        elif line.startswith('**Q'):
            # Extract question
            match = re.match(r'\*\*(Q\d+:.*?)\*\*', line)
            if match:
                p = doc.add_paragraph()
                p.add_run(match.group(1)).bold = True
        
        # Bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            p = doc.add_paragraph(line.strip()[2:], style='List Bullet')
        
        # Regular paragraph
        elif line.strip():
            # Handle inline bold (**text**)
            parts = re.split(r'(\*\*.*?\*\*)', line)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                elif part:
                    p.add_run(part)
        
        i += 1
    
    # Save document
    doc.save("Chapter13_Data_Migration.docx")
    print("✅ Word document created: Chapter13_Data_Migration.docx")

if __name__ == "__main__":
    main()
