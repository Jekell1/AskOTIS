"""
Export Chapter 12 markdown to Word document format.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def parse_markdown_to_word(md_file, docx_file):
    """Convert markdown to Word document with formatting"""
    
    with open(md_file, "r", encoding="utf-8") as f:
        content = f.read()
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 128)
                code_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Chapter title (# Title)
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(24)
                run.font.bold = True
        
        # Major sections (## Section)
        elif line.startswith('## '):
            section = line[3:].strip()
            p = doc.add_heading(section, level=1)
            for run in p.runs:
                run.font.size = Pt(16)
                run.font.bold = True
        
        # Subsections (### Subsection)
        elif line.startswith('### '):
            subsection = line[4:].strip()
            p = doc.add_heading(subsection, level=2)
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.bold = True
        
        # Questions (bold text starting with **)
        elif line.startswith('**') and line.endswith('**'):
            question = line[2:-2].strip()
            p = doc.add_paragraph()
            run = p.add_run(question)
            run.bold = True
            run.font.size = Pt(11)
        
        # Bullet points
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
        
        # Numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
        
        # Regular paragraphs
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line.strip())
        
        i += 1
    
    doc.save(docx_file)
    print(f"✅ Word document created: {docx_file}")

def add_formatted_text(paragraph, text):
    """Add text to paragraph with bold formatting support"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            # Regular text
            paragraph.add_run(part)

def main():
    md_file = "Chapter12_Building_Frontend.md"
    docx_file = "Chapter12_Building_Frontend.docx"
    
    print(f"Converting {md_file} to Word format...")
    parse_markdown_to_word(md_file, docx_file)

if __name__ == "__main__":
    main()
