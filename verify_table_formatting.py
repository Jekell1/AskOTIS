#!/usr/bin/env python3
"""
Verify that Word documents have proper table formatting (not text tables)
"""

from docx import Document
import sys

def check_tables_in_document(filename):
    """Check if a Word document contains properly formatted tables."""
    try:
        doc = Document(filename)
        table_count = len(doc.tables)
        
        print(f"\n{'='*70}")
        print(f"Document: {filename}")
        print(f"{'='*70}")
        print(f"Number of tables found: {table_count}")
        
        if table_count == 0:
            print("  No tables found in this document.")
            return False
        
        # Check each table
        for i, table in enumerate(doc.tables, 1):
            print(f"\nTable {i}:")
            print(f"  Rows: {len(table.rows)}")
            print(f"  Columns: {len(table.columns)}")
            print(f"  Style: {table.style.name if table.style else 'None'}")
            
            # Show first row (header)
            if len(table.rows) > 0:
                header_cells = table.rows[0].cells
                header_text = [cell.text[:30] + '...' if len(cell.text) > 30 else cell.text 
                              for cell in header_cells]
                print(f"  Header: {' | '.join(header_text)}")
                
                # Check if header is bold
                is_bold = False
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.bold:
                                is_bold = True
                                break
                print(f"  Header formatting: {'Bold ✓' if is_bold else 'Not bold ✗'}")
        
        return True
        
    except Exception as e:
        print(f"Error reading {filename}: {e}")
        return False

def main():
    # Check key chapters
    test_files = [
        'Chapter1_Introduction.docx',
        'Chapter2_Business_Domain.docx',
        'Chapter5_Modernization_Goals.docx',
    ]
    
    print("\n" + "="*70)
    print("Verifying Table Formatting in Word Documents")
    print("="*70)
    
    found_tables = False
    for filename in test_files:
        if check_tables_in_document(filename):
            found_tables = True
    
    print("\n" + "="*70)
    if found_tables:
        print("SUCCESS: Found properly formatted tables with Word table objects")
        print("(Not plain text with | characters)")
    else:
        print("No tables found in test documents")
    print("="*70)

if __name__ == "__main__":
    main()
